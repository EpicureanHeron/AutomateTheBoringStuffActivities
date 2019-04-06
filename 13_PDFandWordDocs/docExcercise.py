import docx

def ex1():
    doc = docx.Document('demo.docx')
    print(len(doc.paragraphs))

    print(doc.paragraphs[0].text)

    print(doc.paragraphs[1].text)

    print(len(doc.paragraphs[1].runs))

    print(doc.paragraphs[1].runs[0].text)

    print(doc.paragraphs[1].runs[1].text)

    print(doc.paragraphs[1].runs[2].text)

    print(doc.paragraphs[1].runs[3].text)

#ex1()

def ex2(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(ex2('demo.docx'))